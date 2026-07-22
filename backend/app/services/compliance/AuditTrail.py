from io import BytesIO

from app.schemas.compliance import ComplianceReport


class AuditTrail:
    """
    Handles formatting and binary byte export for PDF, DOCX, and JSON compliance reports.
    """

    def export_pdf(self, report: ComplianceReport) -> bytes:
        from reportlab.lib.colors import HexColor
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40
        )

        styles = getSampleStyleSheet()

        # Premium dark navy title style
        title_style = ParagraphStyle(
            "ReportTitle",
            parent=styles["Heading1"],
            fontSize=22,
            leading=26,
            textColor=HexColor("#0F172A"),
            spaceAfter=15,
        )

        section_heading = ParagraphStyle(
            "SectionHeading",
            parent=styles["Heading2"],
            fontSize=14,
            leading=18,
            textColor=HexColor("#0EA5E9"),
            spaceBefore=12,
            spaceAfter=6,
        )

        body_style = ParagraphStyle(
            "BodyText",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            textColor=HexColor("#334155"),
            spaceAfter=8,
        )

        story = []

        story.append(Paragraph("APEX INDUSTRIAL COMPLIANCE REPORT", title_style))
        story.append(Paragraph(f"<b>Report ID:</b> {report.report_id}", body_style))
        story.append(Spacer(1, 10))

        story.append(Paragraph("Incident Summary", section_heading))
        story.append(Paragraph(report.incident_summary, body_style))

        story.append(Paragraph("Root Cause Analysis", section_heading))
        story.append(Paragraph(report.root_cause, body_style))

        story.append(Paragraph("Final Resolution Outcome", section_heading))
        story.append(Paragraph(report.final_resolution, body_style))

        story.append(Paragraph("Technician Actions", section_heading))
        for act in report.technician_actions:
            story.append(Paragraph(f"• {act}", body_style))

        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        return pdf_bytes

    def export_docx(self, report: ComplianceReport) -> bytes:
        from docx import Document

        doc = Document()
        doc.add_heading("APEX INDUSTRIAL COMPLIANCE REPORT", level=0)

        doc.add_paragraph(f"Report ID: {report.report_id}")

        doc.add_heading("Incident Summary", level=1)
        doc.add_paragraph(report.incident_summary)

        doc.add_heading("Root Cause Analysis", level=1)
        doc.add_paragraph(report.root_cause)

        doc.add_heading("Final Resolution Outcome", level=1)
        doc.add_paragraph(report.final_resolution)

        doc.add_heading("Technician Actions", level=1)
        for act in report.technician_actions:
            doc.add_paragraph(act, style="List Bullet")

        buffer = BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        return docx_bytes
